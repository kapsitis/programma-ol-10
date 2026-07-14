#!/usr/bin/env python3
"""Convert a Markdown file to a styled .docx document.

This is a "template"-style converter for Word, analogous to what the Eisvogel
template does for PDF output. It works in two stages:

  1. pandoc converts the Markdown body to .docx, turning ``$$...$$`` math into
     native Word (OMML) equations. Paragraph/character styles and the base
     look come from an optional pandoc *reference document* (--reference-doc),
     the DOCX equivalent of a LaTeX template.

  2. python-docx post-processes the resulting file to add things pandoc cannot
     express from front-matter: a running header, a (multi-line) running
     footer, and a base font/size. Pandoc has no front-matter mechanism for
     running headers/footers, so this second pass is what makes them possible.

Header/footer/font are read from the Markdown YAML front-matter so they travel
with the document:

    ---
    docx_header: "1. temats. ..."
    docx_footer:
      - "© Valsts izglītības attīstības aģentūra | ESF+ projekts Nr. ..."
      - "Pedagogu profesionālā atbalsta sistēmas izveide"
    docx_font: "Calibri"
    docx_fontsize: 11
    ---

Any of these may be omitted (or overridden on the command line). ``docx_footer``
may be a single string or a list of strings (one paragraph per line).

Usage:
    python scripts/md_to_docx.py <md-file> <docx-file> [options]

Requirements: pandoc on PATH, and the Python packages python-docx and pyyaml.
"""
from __future__ import annotations

import argparse
import re
import shutil
import subprocess
import sys
import tempfile
from pathlib import Path

# The Windows console defaults to a legacy code page (cp1252) that cannot encode
# Latvian diacritics; make our own stdout/stderr UTF-8 so status output works.
for _stream in (sys.stdout, sys.stderr):
    try:
        _stream.reconfigure(encoding="utf-8")
    except (AttributeError, ValueError):
        pass

try:
    import yaml
except ImportError:
    sys.exit("Missing dependency: PyYAML. Install with: pip install pyyaml")

try:
    import docx
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ImportError:
    sys.exit("Missing dependency: python-docx. Install with: pip install python-docx")


def parse_front_matter(md_path: Path) -> dict:
    """Return the YAML front-matter of a Markdown file as a dict (or {})."""
    text = md_path.read_text(encoding="utf-8")
    if not text.startswith("---"):
        return {}
    # Split on the closing '---' fence of the front-matter block.
    parts = text.split("\n---", 1)
    if len(parts) < 2:
        return {}
    header = parts[0][len("---"):]
    try:
        data = yaml.safe_load(header)
    except yaml.YAMLError as exc:
        sys.exit(f"Could not parse front-matter in {md_path}: {exc}")
    return data if isinstance(data, dict) else {}


_DISPLAY_MATH = re.compile(r"\$\$(.*?)\$\$", re.DOTALL)

# Invisible sentinels (private-use characters) that mark the bounds of an
# explicit ``::: solution`` block. They ride through pandoc as raw paragraphs
# and are consumed (and removed) by apply_solution_blocks() in post-processing.
_SOL_MARK_START = "SOLSTART"
_SOL_MARK_END = "SOLEND"
_SOL_RAW_START = f'\n\n```{{=openxml}}\n<w:p><w:r><w:t>{_SOL_MARK_START}</w:t></w:r></w:p>\n```\n\n'
_SOL_RAW_END = f'\n\n```{{=openxml}}\n<w:p><w:r><w:t>{_SOL_MARK_END}</w:t></w:r></w:p>\n```\n\n'
_SOL_BLOCK = re.compile(r"(?ms)^:::+[ \t]*solution[ \t]*$\n(.*?)^:::+[ \t]*$")


def preprocess_markdown(text: str) -> str:
    """Rewrite the Markdown so solution regions get highlighted in the .docx.

    Two authoring conventions are supported:

      * an explicit ``::: solution`` / ``:::`` fenced block -- the WHOLE region
        (equations, prose, tables, images) is highlighted uniformly. This is the
        reliable choice when a solution contains anything other than equations.
      * any display equation ``$$ ... $$`` that contains ``\\color`` is treated
        as a *solution* and wrapped automatically -- convenient for the common
        case where the whole answer is a single blue formula.
    """
    # 1) Bracket explicit ::: solution ... ::: blocks with invisible sentinels;
    #    post-processing shades every paragraph and table cell between them.
    text = _SOL_BLOCK.sub(
        lambda m: f"{_SOL_RAW_START}{m.group(1)}{_SOL_RAW_END}", text)

    # 2) Auto-wrap \color-bearing display equations as single-paragraph solutions.
    def _wrap(match: "re.Match") -> str:
        body = match.group(1)
        if "\\color" not in body:
            return match.group(0)
        return ('\n\n::: {custom-style="Solution"}\n'
                f"$${body}$$\n:::\n\n")

    return _DISPLAY_MATH.sub(_wrap, text)


def run_pandoc(src_text: str, resource_dir: Path, docx_path: Path,
               reference_doc: Path | None, title: str | None) -> None:
    """Convert preprocessed Markdown text to .docx with pandoc (native math).

    ``src_text`` is the (already preprocessed) Markdown; it is written to a temp
    file inside ``resource_dir`` so relative image paths keep resolving.
    ``title`` overrides the front-matter ``title``: pass an empty string to
    suppress pandoc's visible Title paragraph while leaving the Markdown's own
    ``title`` (used by Jekyll) untouched.
    """
    pandoc = shutil.which("pandoc")
    if not pandoc:
        sys.exit("pandoc not found on PATH. Install pandoc: https://pandoc.org/installing.html")

    tmp = tempfile.NamedTemporaryFile(
        "w", encoding="utf-8", suffix=".md", dir=str(resource_dir), delete=False)
    try:
        tmp.write(src_text)
        tmp.close()
        cmd = [
            pandoc,
            tmp.name,
            "-o", str(docx_path),
            "--from", "markdown",
            # Resolve images (![](....png)) relative to the Markdown folder.
            "--resource-path", str(resource_dir),
        ]
        if title is not None:
            # Command-line metadata overrides the front-matter value; an empty
            # string makes pandoc's $if(title)$ template test false -> no Title.
            cmd += ["--metadata", f"title={title}"]
        if reference_doc:
            if not reference_doc.exists():
                sys.exit(f"Reference doc not found: {reference_doc}")
            cmd += ["--reference-doc", str(reference_doc)]

        result = subprocess.run(cmd, capture_output=True, text=True)
        if result.returncode != 0:
            sys.exit(f"pandoc failed:\n{result.stderr}")
    finally:
        Path(tmp.name).unlink(missing_ok=True)


def set_base_font(document: "docx.document.Document", font_name: str | None,
                  font_size: float | None) -> None:
    """Set the document's default (Normal style) font name and/or size."""
    if not font_name and not font_size:
        return
    normal = document.styles["Normal"]
    if font_name:
        normal.font.name = font_name
        # Ensure the name also applies to complex-script / east-asian runs.
        rpr = normal.element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = rpr.makeelement(qn("w:rFonts"), {})
            rpr.append(rfonts)
        for attr in ("w:ascii", "w:hAnsi", "w:cs"):
            rfonts.set(qn(attr), font_name)
    if font_size:
        normal.font.size = Pt(float(font_size))


def _get_or_add_style(document, name: str):
    """Return the paragraph style ``name``, creating it if pandoc didn't."""
    try:
        return document.styles[name]
    except KeyError:
        return document.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)


def ensure_semantic_styles(document, solution_fill: str) -> None:
    """Define the Solution paragraph style used for worked answers.

    Solution = a light background shading only (no border/accent bar), so the
    answer reads as a distinct block without heavy emphasis. Problem statements
    are left as plain text, so no Problem style is needed.

    ``solution_fill`` is a hex RGB string (no '#'). The style is created if
    pandoc left it undefined, or updated in place if it referenced it already.
    """
    sol = _get_or_add_style(document, "Solution")
    ppr = sol.element.get_or_add_pPr()
    # Clear any previously applied shading/border, then apply shading only.
    for tag in ("w:shd", "w:pBdr"):
        existing = ppr.find(qn(tag))
        if existing is not None:
            ppr.remove(existing)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), solution_fill)
    ppr.append(shd)


def _shade(element, fill: str) -> None:
    """Give a <w:pPr> or <w:tcPr> a solid ``fill`` background (replaces any)."""
    old = element.find(qn("w:shd"))
    if old is not None:
        element.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    element.append(shd)


def _set_para_style(p, name: str) -> None:
    """Force paragraph ``p`` to use the named paragraph style."""
    ppr = p.find(qn("w:pPr"))
    if ppr is None:
        ppr = OxmlElement("w:pPr")
        p.insert(0, ppr)
    style = ppr.find(qn("w:pStyle"))
    if style is None:
        style = OxmlElement("w:pStyle")
        ppr.insert(0, style)
    style.set(qn("w:val"), name)


def apply_solution_blocks(document, solution_fill: str) -> None:
    """Uniformly highlight everything inside an explicit ``::: solution`` block.

    The block is delimited by the sentinel paragraphs injected in
    preprocess_markdown(). Every paragraph between them is given the Solution
    style (shaded), every table cell is filled with the same colour, and the
    sentinels themselves are deleted. This keeps prose and tables -- which
    pandoc's paragraph style cannot reach -- on the same background as the math.
    """
    body = document.element.body
    w_p, w_tbl, w_tc = qn("w:p"), qn("w:tbl"), qn("w:tc")
    inside = False
    sentinels = []
    for child in list(body):
        if child.tag == w_p:
            txt = "".join(t.text or "" for t in child.iter(qn("w:t"))).strip()
            if txt == _SOL_MARK_START:
                inside, _ = True, sentinels.append(child)
                continue
            if txt == _SOL_MARK_END:
                inside, _ = False, sentinels.append(child)
                continue
            if inside:
                _set_para_style(child, "Solution")
        elif child.tag == w_tbl and inside:
            for tc in child.iter(w_tc):
                tcpr = tc.find(qn("w:tcPr"))
                if tcpr is None:
                    tcpr = OxmlElement("w:tcPr")
                    tc.insert(0, tcpr)
                _shade(tcpr, solution_fill)
    for child in sentinels:
        body.remove(child)


def set_math_alignment(document, alignment: str | None) -> None:
    """Set the justification of every display equation (``$$...$$``).

    Pandoc emits each display equation as an <m:oMathPara> whose <m:jc> defaults
    to "center". Word keeps display-style (full-size fractions etc.) regardless
    of justification, so switching this to "left" gives left-aligned, still
    displaystyle math -- which plain ``$$...$$`` cannot express directly.
    Valid values: left, center, right (None leaves pandoc's default).
    """
    if not alignment or alignment == "center":
        return
    m_uri = "http://schemas.openxmlformats.org/officeDocument/2006/math"
    jc = f"{{{m_uri}}}jc"
    para_pr = f"{{{m_uri}}}oMathParaPr"
    math_para = f"{{{m_uri}}}oMathPara"
    val = f"{{{m_uri}}}val"
    body = document.element.body
    for opara in body.iter(math_para):
        ppr = opara.find(para_pr)
        if ppr is None:
            ppr = opara.makeelement(para_pr, {})
            opara.insert(0, ppr)
        el = ppr.find(jc)
        if el is None:
            el = ppr.makeelement(jc, {})
            ppr.append(el)
        el.set(val, alignment)


def render_horizontal_rules(document) -> None:
    """Turn pandoc's thematic breaks (``---``) into visible ruled lines.

    Pandoc renders ``---`` as a legacy VML horizontal rule (<v:rect o:hr="t">
    with width:0), which Word shows but LibreOffice/other viewers may not. We
    replace each such paragraph with an empty paragraph carrying a real bottom
    border, giving a portable, full-width writing line (e.g. a name/class line).
    """
    w = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    v_rect = "{urn:schemas-microsoft-com:vml}rect"
    body = document.element.body
    for para in list(body.iter(f"{{{w}}}p")):
        if para.find(f".//{v_rect}") is None:
            continue
        # Drop existing runs (the VML rule) and give the paragraph a bottom border.
        for run in para.findall(f"{{{w}}}r"):
            para.remove(run)
        ppr = para.find(f"{{{w}}}pPr")
        if ppr is None:
            ppr = para.makeelement(f"{{{w}}}pPr", {})
            para.insert(0, ppr)
        pbdr = para.makeelement(f"{{{w}}}pBdr", {})
        bottom = para.makeelement(f"{{{w}}}bottom", {
            f"{{{w}}}val": "single",
            f"{{{w}}}sz": "8",        # 8 eighths of a point = 1.0 pt
            f"{{{w}}}space": "1",
            f"{{{w}}}color": "000000",
        })
        pbdr.append(bottom)
        ppr.append(pbdr)


_PAPER_SIZES_CM = {  # (width, height)
    "a4": (21.0, 29.7), "a5": (14.8, 21.0),
    "letter": (21.59, 27.94), "legal": (21.59, 35.56),
}


def apply_page_geometry(document, geometry: str | None) -> None:
    """Set page size and margins on every section from a LaTeX-style geometry.

    Pandoc ignores the ``geometry:`` front-matter for docx, so the output has no
    explicit page size (Word then assumes its default). We mirror the declared
    geometry (e.g. ``a4paper, top=2.5cm, left=2cm``) so the .docx matches the
    intended A4 layout -- which also gives the header its exact content width.
    """
    if not geometry:
        return
    g = geometry.lower()
    dims = next((v for k, v in _PAPER_SIZES_CM.items() if k in g), None)
    margins = {m.group(1): float(m.group(2))
               for m in re.finditer(r"(top|bottom|left|right)\s*=\s*([\d.]+)\s*cm", g)}
    for section in document.sections:
        if dims:
            section.page_width, section.page_height = Cm(dims[0]), Cm(dims[1])
        for edge, value in margins.items():
            setattr(section, f"{edge}_margin", Cm(value))


def _content_width(section):
    """Text-column width of a section, with an A4/2cm fallback if unset."""
    if section.page_width and section.left_margin is not None \
            and section.right_margin is not None:
        return section.page_width - section.left_margin - section.right_margin
    return Cm(17.0)  # A4 (21 cm) minus 2 cm margins each side


def _style_run(run, font_name: str | None) -> None:
    """Apply the shared small-grey header/footer look to a run."""
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    if font_name:
        run.font.name = font_name


def _fill_frame(frame, lines: list[str], font_name: str | None,
                alignment) -> None:
    """Write text lines into a header/footer as a SINGLE paragraph.

    Multiple lines are separated by a forced line break (<w:br/>) within one
    run, so the whole block stays one paragraph rather than several.
    """
    frame.is_linked_to_previous = False
    para = frame.paragraphs[0]  # reuse the frame's existing empty paragraph
    para.alignment = alignment
    run = para.add_run()
    run.text = "\n".join(lines)  # '\n' is rendered as a <w:br/> by python-docx
    _style_run(run, font_name)


def _fill_header(document, section, left: str, right: str | None,
                 font_name: str | None) -> None:
    """Header with ``left`` at the left margin and ``right`` at the right margin.

    A single right-aligned tab stop is placed at the content width and a tab
    separates the two texts, so they are pushed to the opposite page edges. To
    keep Word's built-in Header tab stops (a centre tab) from catching the tab
    first, the Header style is redefined with no tabs of its own.
    """
    frame = section.header
    frame.is_linked_to_previous = False
    para = frame.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if right:
        # Neutralise the built-in Header style's default centre/right tabs.
        _clear_style_tabs(_get_or_add_style(document, "Header"))
        para.paragraph_format.tab_stops.add_tab_stop(
            _content_width(section), WD_TAB_ALIGNMENT.RIGHT)
        run = para.add_run()
        run.text = f"{left}\t{right}"
    else:
        run = para.add_run()
        run.text = left
    _style_run(run, font_name)


def _clear_style_tabs(style) -> None:
    """Remove any <w:tabs> from a paragraph style's properties."""
    ppr = style.element.find(qn("w:pPr"))
    if ppr is not None:
        tabs = ppr.find(qn("w:tabs"))
        if tabs is not None:
            ppr.remove(tabs)


def apply_header_footer(document, header_left, header_right, footer_lines,
                        font_name) -> None:
    """Set a running header (left + right) and a right-aligned footer."""
    for section in document.sections:
        if header_left or header_right:
            _fill_header(document, section, header_left or "", header_right, font_name)
        if footer_lines:
            _fill_frame(section.footer, footer_lines, font_name,
                        WD_ALIGN_PARAGRAPH.RIGHT)


def normalize_footer(value) -> list[str]:
    """Accept a string or list for the footer; return a list of lines."""
    if value is None:
        return []
    if isinstance(value, str):
        return [value]
    return [str(v) for v in value]


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Convert a Markdown file to a styled .docx (running header/footer, base font)."
    )
    parser.add_argument("md_file", help="Input Markdown file")
    parser.add_argument("docx_file", help="Output .docx file")
    parser.add_argument("--reference-doc", help="Optional pandoc reference .docx for styles")
    parser.add_argument("--header", help="Override the running header (left) text")
    parser.add_argument("--header-right", help="Override the running header (right) text")
    parser.add_argument("--footer", action="append",
                        help="Override a running footer line (repeatable)")
    parser.add_argument("--font", help="Override the base font name")
    parser.add_argument("--fontsize", type=float, help="Override the base font size (pt)")
    parser.add_argument("--title", help="Set the visible document title (default: suppressed)")
    parser.add_argument("--math-align", choices=("left", "center", "right"),
                        help="Justify display equations (default from front-matter or center)")
    args = parser.parse_args()

    md_path = Path(args.md_file)
    docx_path = Path(args.docx_file)
    if not md_path.exists():
        sys.exit(f"Input file not found: {md_path}")
    docx_path.parent.mkdir(parents=True, exist_ok=True)

    fm = parse_front_matter(md_path)
    header_text = args.header if args.header is not None else fm.get("docx_header")
    header_right = (args.header_right if args.header_right is not None
                    else fm.get("docx_header_right"))
    footer_lines = (normalize_footer(args.footer) if args.footer
                    else normalize_footer(fm.get("docx_footer")))
    font_name = args.font or fm.get("docx_font")
    font_size = args.fontsize if args.fontsize is not None else fm.get("docx_fontsize")
    math_align = args.math_align or fm.get("docx_math_align")
    solution_fill = str(fm.get("docx_solution_fill", "DCE6F1"))   # light blue
    # Suppress pandoc's visible Title unless one is explicitly requested. The
    # front-matter `title` is kept for Jekyll; here we blank it for the .docx.
    if args.title is not None:
        title = args.title
    elif fm.get("docx_title"):
        title = str(fm["docx_title"])
    else:
        title = ""  # blank -> pandoc emits no Title paragraph

    src_text = preprocess_markdown(md_path.read_text(encoding="utf-8"))
    ref = Path(args.reference_doc) if args.reference_doc else None
    run_pandoc(src_text, md_path.parent, docx_path, ref, title)

    document = docx.Document(str(docx_path))
    apply_page_geometry(document, fm.get("geometry"))
    set_base_font(document, font_name, font_size)
    ensure_semantic_styles(document, solution_fill)
    apply_solution_blocks(document, solution_fill)
    set_math_alignment(document, math_align)
    render_horizontal_rules(document)
    apply_header_footer(document, header_text, header_right, footer_lines, font_name)
    document.save(str(docx_path))

    print(f"Wrote {docx_path}")
    print(f"  title:  {title!r} (suppressed)" if title == "" else f"  title:  {title}")
    if header_text:
        print(f"  header: {header_text}")
    if header_right:
        print(f"  header (right): {header_right}")
    for line in footer_lines:
        print(f"  footer (right): {line}")
    if font_name or font_size:
        print(f"  font:   {font_name or '(unchanged)'} {font_size or ''}".rstrip())
    if math_align:
        print(f"  math:   {math_align}-aligned")
    print(f"  marks:  solutions shaded #{solution_fill} (no bar); problems plain text")


if __name__ == "__main__":
    main()
